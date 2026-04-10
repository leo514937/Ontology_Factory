from __future__ import annotations

import importlib
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:  # type: ignore[override]
        if tag in {"script", "style", "nav", "footer", "header"}:
            self._skip_depth += 1
            return
        if self._skip_depth == 0 and tag in {"p", "div", "section", "article", "br", "li", "h1", "h2", "h3", "h4"}:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        if tag in {"script", "style", "nav", "footer", "header"} and self._skip_depth > 0:
            self._skip_depth -= 1
            return
        if self._skip_depth == 0 and tag in {"p", "div", "section", "article", "li"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        if self._skip_depth == 0 and data.strip():
            self._parts.append(data)

    def get_text(self) -> str:
        return "\n".join(part.strip() for part in self._parts if part.strip())


def _charset_normalizer_from_bytes():
    try:
        return importlib.import_module("charset_normalizer").from_bytes
    except ModuleNotFoundError:
        return None


def _beautiful_soup():
    try:
        return importlib.import_module("bs4").BeautifulSoup
    except ModuleNotFoundError:
        return None


def _decode_bytes(b: bytes, encoding_fallbacks: Iterable[str]) -> str:
    for enc in encoding_fallbacks:
        try:
            return b.decode(enc)
        except Exception:
            pass
    try:
        from_bytes = _charset_normalizer_from_bytes()
        if from_bytes is not None:
            best = from_bytes(b).best()
            if best is not None:
                return str(best)
    except Exception:
        pass
    return b.decode("utf-8", errors="replace")


def load_document(path: str, encoding_fallbacks: Iterable[str]) -> LoadedDocument:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(str(p))

    suffix = p.suffix.lower()
    if suffix in {".txt", ".md"}:
        return LoadedDocument(path=p, text=_decode_bytes(p.read_bytes(), encoding_fallbacks))

    if suffix in {".html", ".htm"}:
        html = _decode_bytes(p.read_bytes(), encoding_fallbacks)
        soup_factory = _beautiful_soup()
        if soup_factory is not None:
            parser_name = "lxml"
            try:
                importlib.import_module("lxml")
            except ModuleNotFoundError:
                parser_name = "html.parser"
            soup = soup_factory(html, parser_name)
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            text = soup.get_text("\n")
        else:
            parser = _HTMLTextExtractor()
            parser.feed(html)
            text = parser.get_text()
        return LoadedDocument(path=p, text=text)

    if suffix == ".docx":
        from docx import Document

        doc = Document(str(p))
        return LoadedDocument(path=p, text="\n".join(para.text for para in doc.paragraphs))

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(p))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        return LoadedDocument(path=p, text="\n".join(parts))

    raise ValueError(f"Unsupported file type: {suffix}")


def discover_inputs(input_globs: Iterable[str], base_dir: str) -> list[str]:
    base = Path(base_dir)
    paths: set[Path] = set()
    for pattern in input_globs:
        for path in base.glob(pattern):
            if path.is_file():
                paths.add(path.resolve())
    return [str(path) for path in sorted(paths)]


_CONTROL_CHARS = re.compile(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F]")


def normalize_text_for_pipeline(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = _CONTROL_CHARS.sub("", normalized)
    normalized = normalized.replace("\u3000", " ")
    return normalized
